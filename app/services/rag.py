from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain.chains import RetrievalQA
# from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
# from langchain_huggingface import HuggingFaceEmbeddings
from io import BytesIO
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
from langchain.docstore.document import Document
from dotenv import load_dotenv
import os
from db.chroma_store import get_chroma_client

load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

def load_document(file_stream: BytesIO, file_extension: str):
    docs = []

    if file_extension == "txt":
        text = file_stream.read().decode("utf-8", errors="ignore")
        docs.append(Document(page_content=text, metadata={}))

    elif file_extension == "pdf":
        pdf_reader = PdfReader(file_stream)
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        docs.append(Document(page_content=text.strip(), metadata={}))

    elif file_extension == "docx":
        doc = DocxDocument(file_stream)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        docs.append(Document(page_content=text, metadata={}))

    else:
        raise ValueError(f"Unsupported file type: {file_extension}")

    return docs

def split_documents(documents):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )
    return text_splitter.split_documents(documents)

def create_embeddings_and_store(documents, collection_name: str):
    # embeddings = HuggingFaceEmbeddings(
    #     model_name="sentence-transformers/all-MiniLM-L6-v2",
    # )
    embeddings = GoogleGenerativeAIEmbeddings(model="models/text-embedding-004", google_api_key=GEMINI_API_KEY)
    # Use a persistent client for ChromaDB
    client = get_chroma_client()
    # client = chromadb.PersistentClient(path="../chroma_db")
    vector_store = Chroma.from_documents(
        documents=documents,
        embedding=embeddings,
        client=client,
        collection_name=collection_name,
    )
    return vector_store

def get_rag_chain(collection_name: str):
    # embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    embeddings = GoogleGenerativeAIEmbeddings(model="models/text-embedding-004", google_api_key=GEMINI_API_KEY)
    client = get_chroma_client()
    vector_store = Chroma(client=client, collection_name=collection_name, embedding_function=embeddings)
    
    llm = ChatGoogleGenerativeAI(model="gemini-3-flash-preview", temperature=0.5, google_api_key=GEMINI_API_KEY)
    qa_chain = RetrievalQA.from_chain_type(
        llm=llm,
        chain_type="stuff",
        retriever=vector_store.as_retriever()
    )
    # print(qa_chain)
    return qa_chain
